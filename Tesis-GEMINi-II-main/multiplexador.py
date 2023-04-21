#!/usr/bin/env python
# coding: utf-8

# In[1]:


#importar librerias
import os
from os import path
import csv
import primer3
from Bio import SeqIO
from Bio.Seq import Seq
import copy
import docx
import networkx as nx

#globales
gmvc=50


# In[2]:


#definir funciones
def calcTemp(dna):
    return sum([4 if n in ["G","C"] else 2 for n in dna.upper()])

def calcGC(dna):
    return sum(1 for n in dna.upper() if n in ["G","C"])/len(dna)

def checkComp(pp1,pp2,k,e,mindg,maxdt,mindlen):
    """!
    Función que se encarga de revisar compatibilidad entre dos pares de partidores. Almacena las causas de incompatibilidad en un string (rson).
    \param <pp1> {Diccionario, tiene los datos de un par de partidores.}
    \param <pp2> {Diccionario, tiene los datos de un par de partidores.}
    \param <k> {Entero, posición asociada al primer par de partidores.}
    \param <e> {Entero, posición asociada al segundo par de partidores.}
    \param <mindg> {Numero, minimo deltaG tolerado.}
    \param <maxdt> {Numero, maxima diferencia de temperatura tolerada.}
    \param <mindlen> {Entero, minima diferencia de largo requerida.}
    """ 
    #print("temp es "+str(gmvc))
    rson=[]
    comp=True
    f1=pp1["f"]
    r1=pp1["r"]
    f2=pp2["f"]
    r2=pp2["r"]
    gibf1f2=primer3.calcHeterodimer(f1,f2,mv_conc=gmvc).dg
    gibf1r2=primer3.calcHeterodimer(f1,r2,mv_conc=gmvc).dg
    gibr1f2=primer3.calcHeterodimer(r1,f2,mv_conc=gmvc).dg
    gibr1r2=primer3.calcHeterodimer(r1,r2,mv_conc=gmvc).dg
    dt=abs(pp1["tp"]-pp2["tp"])
    dsize=abs(pp1["ampsize"]-pp2["ampsize"])
    rpe=pp2["relpos"]
    rpk=pp1["relpos"]
    if abs(rpe-rpk)==1:
        comp=False
        rson.append("Contiguedad")
    if gibf1f2<mindg:
        comp=False
        rson.append("F{} y F{} dan Gibbs {}".format(k,e,gibf1f2))
    if gibf1r2<mindg:
        comp=False
        rson.append("F{} y R{} dan Gibbs {}".format(k,e,gibf1r2))
    if gibr1f2<mindg:
        comp=False
        rson.append("R{} y F{} dan Gibbs {}".format(k,e,gibr1f2))
    if gibr1r2<mindg:
        comp=False
        rson.append("R{} y R{} dan Gibbs {}".format(k,e,gibr1r2))
    if dt>maxdt:
        comp=False
        rson.append("Diferencia de temperatura es {}".format(dt))
    if dsize<mindlen:
        comp=False
        rson.append("Diferencia de tamaño es {}".format(dsize))
    return rson
    
def testHet(pp1,pp2):
    f1=pp1["f"]
    r1=pp1["r"]
    f2=pp2["f"]
    r2=pp2["r"]
    print("f1f2: "+str(primer3.calcHeterodimer(f1,f2).dg))
    print("f1r2: "+str(primer3.calcHeterodimer(f1,r2).dg))
    print("r1f2: "+str(primer3.calcHeterodimer(r1,f2).dg))
    print("r1r2: "+str(primer3.calcHeterodimer(r1,r2).dg))


# In[3]:


#funciones fasta
def findSec(myseq,tgseq,s):
    """!
    Función que busca y retorna la posición de un partidor en el gen.
    \param <myseq> {Seq, secuencia del gen.}
    \param <tgseq> {String, secuencia de partidor.}
    \param <s> {String, indica si partidor es Forward("F") o Reverse("R").}
    """ 
    if type(tgseq)==str:
        tgseq=Seq(tgseq.upper())
    if s=="L":
        posi=myseq.find(tgseq)
        pose=posi+len(tgseq)
        return posi,pose
    if s=="R":
        tempseq=tgseq.reverse_complement()
        return findSec(myseq,tempseq,"L")

def loadFastaFile(filename):
    seqlist=[]
    for seq_record in SeqIO.parse(filename, "fasta"):
        tempdict={"id":seq_record.id,"seq":seq_record.seq,"record":seq_record}
        seqlist.append(tempdict)
    return seqlist

def Fasta2Seq(filename):
    return loadFastaFile(filename)[0]["seq"]

def extendPrimer(myseq,tgseq,s,l):
    """!
    Función para extender un partidor.
    \param <myseq> {Seq, secuencia del gen.}
    \param <tgseq> {String, secuencia de partidor.}
    \param <s> {String, indica si partidor es Forward("F") o Reverse("R").}
    \param <l> {Entero, cuantos nucleotidos se extenderá el partidor.}
    """ 
    if type(tgseq)==str:
        tgseq=Seq(tgseq.upper())
    pi,pe=findSec(myseq,tgseq,s)
    if s=="L":
        return myseq[pi:pe+l]
    if s=="R":
        return myseq[pi-l:pe].reverse_complement()

def extendTogether(myseq,tgl,tgr,l):
    """!
    Función que busca extender un F y R juntos. Util en casos donde el R no sea único, pues buscará solo los que existan despues del F.
    \param <myseq> {Seq, secuencia del gen.}
    \param <tgl> {String, secuencia de partidor F.}
    \param <tgr> {String, secuencia de partidor R.}
    \param <s> {String, indica si partidor es Forward("F") o Reverse("R").}
    \param <l> {Entero, cuantos nucleotidos se extenderán los partidores.}
    """ 
    pif,pef=findSec(myseq,tgl,"L")
    pir,per=findSec(myseq[pif:],tgr,"R")
    pir+=pif
    per+=pif
    sl=myseq[pif:pef+l]
    sr=myseq[pir-l:per].reverse_complement()
    return sl,sr

def extendLWG(myseq,tgseq,posi,l):
    #if type(tgseq)==str:
    #    tgseq=Seq(tgseq.upper())
    return myseq[posi:posi+len(tgseq)+l]

def extendRWGL(myseq,tgl,pif,tgr,l):
    pef=pif+len(tgl)
    pir,per=findSec(myseq[pif:],tgr,"R")
    pir+=pif
    per+=pif
    sr=myseq[pir-l:per].reverse_complement()
    return sr


# In[4]:


def loadWordFile(filename):
    script_dir = os.path.abspath(__file__)
    last_backslash_index = script_dir.rfind('/')
    genes_dir = script_dir[0:last_backslash_index]+"/Genes/"+filename
    doc_gene = docx.Document(genes_dir)

    #doc_gene = docx.Document(filename)
    full_gene=""
    for paragraph in doc_gene.paragraphs:
        full_gene = full_gene + paragraph.text   
    seqstr=''.join([i for i in full_gene if not (i.isdigit() or i in[" "])])
    return seqstr

def Word2Seq(filename):
    seqstr=loadWordFile(filename)
    wordseq=Seq(seqstr.upper())
    return wordseq


# In[5]:


#importar información
def loadPrimPairs(filename,filetype):
    """!
    Función para cargar paneles de partidores desde archivo.
    \param <filename> {String, nombre del archivo.}
    \param <filetype> {String, formato del archivo (por ejemplo, si es "GEMIN1" o "articnCoV").}
    """ 
    if filetype=="GEMIN1":
        return loadGEMIN1(filename)
    if filetype=="articnCoV":
        return loadARTICnCoV(filename)
    return {}
def loadGEMIN1(filename,sk=0,rpg=0):
    fields = []
    rows = []
    with open(filename, 'r') as csvfile:
        csvreader = csv.reader(csvfile)

        br=False

        for row in csvreader:
            #print(row)
            if br:
                rows.append(row)
            if len(row)>3 and row[3]=="Largo":#Asegurarse de estar en la fila correcta
                fields = row
                br=True

    #fin del with

    primpairs={}
    for i in range(0,len(rows),2):
        f=rows[i][1]
        r=rows[i+1][1]
        #tp=(calcTemp(f)+calcTemp(r))/2
        tp=(primer3.calcTm(f,mv_conc=gmvc)+primer3.calcTm(r,mv_conc=gmvc))/2
        ampsize=int(rows[i][10])
        primpairs[int(sk+1+i/2)]={"f":f,"r":r,"tp":tp, "ampsize": ampsize, "extend": 1, "relpos": int(sk+rpg+1+i/2),
                                  "givstart":(int(rows[i][2])-1),
                                 "file":filename,"ftype":"gemini"}
    return primpairs
def loadARTICnCoV(filename,sk=0,rpg=0):
    fields = []
    rows = []
    primpairs ={}
    with open(filename) as file:
        tsv_file = csv.reader(file, delimiter="\t")
        for line in tsv_file:
            if fields==[]:
                fields=line
            else:
                rows.append(line)

    for i in range(0,len(rows),2):
            f=rows[i][2]
            r=rows[i+1][2]
            tp=(float(rows[i][5])+float(rows[i+1][5]))/2.0
            ampsize=i*-1000
            primpairs[int(sk+1+i/2)]={"f":f,"r":r,"tp":tp, "ampsize": ampsize, "extend": 1, "relpos": int(sk+rpg+1+i/2),
                                 "file":filename,"ftype":"artic"}
    return primpairs


# In[6]:


def loadMultiple(filenames,filetype):
    """!
    Función para cargar paneles multi-gen.
    \param <filename> {Lista de Strings, nombres de los archivos.}
    \param <filetype> {String, formato de los archivos (por ejemplo, "GEMIN1").}
    """ 
    if filetype=="GEMIN1":
        dic={}
        for i,filename in enumerate(filenames):
            dic.update(loadGEMIN1(filename,len(dic),i))
        return dic


# In[7]:


def oldFastaFixData(primdict,myseq):
    primpairs=copy.deepcopy(primdict)
    for k in primpairs:
        pif,pef=findSec(myseq,primpairs[k]["f"],"L")
        pir,per=findSec(myseq,primpairs[k]["r"],"R")
        primpairs[k]["start"]=pif
        primpairs[k]["end"]=per
        primpairs[k]["calcsize"]=per-pif
    return primpairs

def FastaFixData(primdict,myseq):
    """!
    Función que busca partidores de archivo de paneles en el archivo de secuencia, para detectar incongruencias, y agrega esta información al diccionario de partidores.
    \param <primdict> {Diccionario, diccionario con los partidores cargados.}
    \param <myseq> {Seq, secuencia del gen.}
    """ 
    primpairs=copy.deepcopy(primdict)
    for k in primpairs:
        pif,pef=findSec(myseq,primpairs[k]["f"],"L")
        pir,per=findSec(myseq[pif:],primpairs[k]["r"],"R")
        primpairs[k]["start"]=pif
        primpairs[k]["end"]=per+pif
        primpairs[k]["calcsize"]=per
        if (per-primpairs[k]["ampsize"])!=0:
            primpairs[k]["extend"]=0
    return primpairs


# In[8]:


#revisar compatibilidad
def checkGroupComp(primdict,mindg,maxdt,mindlen):
    """!
    Función que se encarga de revisar compatibilidad entre los partidores de un grupo y agrega un vector de compatibilidad al diccionario de partidores.
    \param <primdict> {Diccionario, diccionario con los partidores cargados.}
    \param <mindg> {Numero, minimo deltaG tolerado}
    \param <maxdt> {Numero, maxima diferencia de temperatura tolerada.}
    \param <mindlen> {Entero, minima diferencia de largo requerida.}
    """ 
    primpairs=copy.deepcopy(primdict)
    for k in primpairs:
        primpairs[k]["reasons"]={}
        primpairs[k]["incvector"]=[]
        for e in primpairs:
            if e!=k:
                reasons=checkComp(primpairs[k],primpairs[e],k,e,mindg,maxdt,mindlen)
                #if abs(e-k)==1 or checkComp(primpairs[k],primpairs[e],k,e,mindg,maxdt,mindlen)==False:
                if len(reasons)>0:
                    primpairs[k]["incvector"].append(1)
                    primpairs[k]["reasons"][e]=reasons
                else:
                    primpairs[k]["incvector"].append(0)
            else:
                primpairs[k]["incvector"].append(0)
    return primpairs


# In[9]:


#algoritmo coloreo para grupos
import time
def coloringGroup(primdict,maxcolor,newfilename="paneles_multiplex.csv"):
    primpairs=copy.deepcopy(primdict)
    primsetnum=len(primpairs)
    incedges=[]
    for k in primpairs:
        incedges.append(primpairs[k]["incvector"])

    pg = Graph(primsetnum)
    pg.graph = incedges

    ncolor=0
    sol=[]
    for i in range (2,maxcolor+1):
        start_time = time.time()
        sol=pg.graphColoring(i)
        #print("--- %s seconds ---" % (time.time() - start_time))
        print("Para {} frascos, coloreo tardo {} segundos".format(i,time.time() - start_time))
        if len(sol)==0:
            print("No se encontro solución para {} colores.".format(i))
        else:
            ncolor=i
            break

    if len(sol)==0:
        print("No se encontraron soluciones para los parámetros fijados.")
    else:
        print("\nSe encontro solución para {} colores:".format(ncolor))
        for cn in range (1,ncolor+1):
            colormems=[i+1 for i, x in enumerate(sol) if x == cn]
            print("Sets de partidores para color {}: ".format(cn) + ', '.join(str(field) for field in colormems))

        #crear archivo csv
        primdictlist=[]
        for k in primpairs:
            setnum=k
            fw=primpairs[k]["f"]
            rv=primpairs[k]["r"]
            tm=primpairs[k]["tp"]
            grn=sol[k-1]
            asp=primpairs[k]["ampsize"]
            tempdict={"SetNumber": setnum, "Forward": fw, "Reverse": rv, "Tm": tm, 'AmpliconSize': asp, "GroupNumber": grn}
            primdictlist.append(tempdict)

        dictfields = ['SetNumber', 'Forward', 'Reverse', 'Tm', 'AmpliconSize', 'GroupNumber']

        with open(newfilename, 'w') as csvfile: 
            writer = csv.DictWriter(csvfile, fieldnames = dictfields) 
            writer.writeheader() 
            writer.writerows(primdictlist) 
            print("\nArchivo {} creado en directorio.".format(newfilename))

##########################################################################################################

#algoritmo coloreo para grupos y returna dict
def coloringGroupDict(primdict,maxcolor,newfilename="paneles_multiplex.csv"):
    primpairs=copy.deepcopy(primdict)
    primsetnum=len(primpairs)
    incedges=[]
    for k in primpairs:
        incedges.append(primpairs[k]["incvector"])

    pg = Graph(primsetnum)
    pg.graph = incedges
    
    dictSol={}
    ncolor=0
    sol=[]
    for i in range (2,maxcolor+1):
        start_time = time.time()
        sol=pg.graphColoring(i)
        #print("--- %s seconds ---" % (time.time() - start_time))
        print("Para {} frascos, coloreo tardo {} segundos".format(i,time.time() - start_time))
        if len(sol)==0:
            print("No se encontro solución para {} colores.".format(i))
        else:
            ncolor=i
            break

    if len(sol)==0:
        print("No se encontraron soluciones para los parámetros fijados.")
    else:
        print("\nSe encontro solución para {} colores:".format(ncolor))
        for cn in range (1,ncolor+1):
            colormems=[i+1 for i, x in enumerate(sol) if x == cn]
            print("Sets de partidores para color {}: ".format(cn) + ', '.join(str(field) for field in colormems))
            dictSol[cn]={"pairs":colormems}
        #crear archivo csv
        primdictlist=[]
        for k in primpairs:
            setnum=k
            fw=primpairs[k]["f"]
            rv=primpairs[k]["r"]
            tm=primpairs[k]["tp"]
            grn=sol[k-1]
            asp=primpairs[k]["ampsize"]
            tempdict={"SetNumber": setnum, "Forward": fw, "Reverse": rv, "Tm": tm, 'AmpliconSize': asp, "GroupNumber": grn}
            primdictlist.append(tempdict)

        dictfields = ['SetNumber', 'Forward', 'Reverse', 'Tm', 'AmpliconSize', 'GroupNumber']

        with open(newfilename, 'w') as csvfile: 
            writer = csv.DictWriter(csvfile, fieldnames = dictfields) 
            writer.writeheader() 
            writer.writerows(primdictlist) 
            print("\nArchivo {} creado en directorio.".format(newfilename))
        return dictSol


# In[10]:


def writePrimFile(primdict,colorset,newfilename="paneles_multiplex.csv"):
    primpairs=copy.deepcopy(primdict)
    colors=copy.deepcopy(colorset)
    for c in colors:
        for t in colors[c]["pairs"]:
            primpairs[t]["color"]=c
    
    primdictlist=[]
    for k in primpairs:
        setnum=k
        fw=primpairs[k]["f"]
        rv=primpairs[k]["r"]
        tm=primpairs[k]["tp"]
        grn=primpairs[k]["color"]
        asp=primpairs[k]["ampsize"]
        tempdict={"SetNumber": setnum, "Forward": fw, "Reverse": rv, "Tm": tm, 'AmpliconSize': asp, "GroupNumber": grn}
        primdictlist.append(tempdict)

    dictfields = ['SetNumber', 'Forward', 'Reverse', 'Tm', 'AmpliconSize', 'GroupNumber']

    with open(newfilename, 'w') as csvfile: 
        writer = csv.DictWriter(csvfile, fieldnames = dictfields) 
        writer.writeheader() 
        writer.writerows(primdictlist) 
        print("\nArchivo {} creado en directorio.".format(newfilename))            


# In[11]:


def tempColors(colors,primers):
    primpairs=copy.deepcopy(primers)
    colorpairs=copy.deepcopy(colors)
    for k in colorpairs:
        temptemps=[]
        for s in colorpairs[k]["pairs"]:
            f=primpairs[s]["f"]
            r=primpairs[s]["r"]      
            temptemps.append(primer3.calcTm(f,mv_conc=gmvc))
            temptemps.append(primer3.calcTm(r,mv_conc=gmvc))
        colorpairs[k]["maxt"]=max(temptemps)
        colorpairs[k]["mint"]=min(temptemps)
    return colorpairs


# In[12]:


def writeReasons(primdict,filename="incomp_reason.txt"):
    """!
    Función lee las razones de incompatibilidad entre partidores y las escribe en un archivo.
    \param <primdict> {Diccionario, diccionario con los partidores cargados (incluyendo vector de compatibilidad).}
    \param <filename> {String, nombre del archivo a escribir.}
    """ 
    primpairs=copy.deepcopy(primdict)
    allstr=""
    for k in primpairs:
        rstr="Incompatibilidades de {}:\n".format(k)
        tdic=primpairs[k]["reasons"]
        for n in tdic:
            rstr+="\tCon {} por:\n".format(n)
            for r in tdic[n]:
                rstr+="\t\t{}\n".format(r)
        #rstr=rstr+tstr
        allstr=allstr+rstr
    with open(filename, 'w') as f:
        f.write(allstr)
        print("\nArchivo {} creado en directorio.".format(filename))
    
#writeReasons(comppairs)


# In[13]:


def findMaxTemp(primers,genseq,maxlen):
    """!
    Función encargada de buscar la temperatura máxima de cada partidor y agregar esta información al diccionario de partidores.
    \param <primers> {Diccionario, diccionario con los partidores.}
    \param <genseq> {Seq, secuencia del gen.}
    \param <maxlen> {Entero, maximo largo que puede tener un partidor.}
    """ 
    primpairs=copy.deepcopy(primers)
    for k in primpairs:
        exte=primpairs[k]["extend"]
        ftype=primpairs[k]["ftype"]
        #print(primpairs[k])
        #print("--------"+str(ext)+"----------")
        if (exte==0 and ftype!="gemini"):
            primpairs[k]["matm"]=primpairs[k]["tp"]
            #print(k,"excluido")
            continue
            
        f=primpairs[k]["f"]
        r=primpairs[k]["r"]
        if (exte==0):
            posil=primpairs[k]["givstart"]
        if maxlen-len(f) > 0:
            if (exte==0):
                fe=extendLWG(genseq,f,posil,maxlen-len(f))
            else:
                fe=extendPrimer(genseq,f,"L",maxlen-len(f))
        else:
            fe=f
        if maxlen-len(r) > 0:
            if (exte==0):
                re=extendRWGL(genseq,f,posil,r,maxlen-len(r))
            else:
                _,re=extendTogether(genseq,f,r,maxlen-len(r))
        else:
            re=r
        
        f=str(fe)
        r=str(re)
        tf=primer3.calcTm(f,mv_conc=gmvc)
        tr=primer3.calcTm(r,mv_conc=gmvc)
        while abs(tf-tr)>1:
            if tf<tr:
                r=r[:-1]
            else:
                f=f[:-1]
            tf=primer3.calcTm(f,mv_conc=gmvc)
            tr=primer3.calcTm(r,mv_conc=gmvc)
        fe=f
        re=r
        #print(k)
        #print(fe,f,tf)
        #print(re,r,tr)
            
        primpairs[k]["matm"]=(primer3.calcTm(str(fe),mv_conc=gmvc)+primer3.calcTm(str(re),mv_conc=gmvc))/2
        primpairs[k]["maxrlen"]=len(re)
        primpairs[k]["maxflen"]=len(fe)
    return primpairs


# In[14]:


#revisar compatibilidad NEO
def checkGroupCompNeo(primdict,mindg,maxdt,mindlen):
    """!
    Función que se encarga de revisar compatibilidad entre los partidores de un grupo y agrega un vector de compatibilidad al diccionario de partidores. Esta versión utiliza rangos de temperatura en vez de temperaturas discretas.
    \param <primdict> {Diccionario, diccionario con los partidores cargados.}
    \param <mindg> {Numero, minimo deltaG tolerado.}
    \param <maxdt> {Numero, maxima diferencia de temperatura tolerada.}
    \param <mindlen> {Entero, minima diferencia de largo requerida.}
    """ 
    primpairs=copy.deepcopy(primdict)
    for k in primpairs:
        primpairs[k]["reasons"]={}
        primpairs[k]["incvector"]=[]
        for e in primpairs:
            if e!=k:
                reasons=checkCompNeo(primpairs[k],primpairs[e],k,e,mindg,maxdt,mindlen)
                #if abs(e-k)==1 or checkComp(primpairs[k],primpairs[e],k,e,mindg,maxdt,mindlen)==False:
                if len(reasons)>0:
                    primpairs[k]["incvector"].append(1)
                    primpairs[k]["reasons"][e]=reasons
                else:
                    primpairs[k]["incvector"].append(0)
            else:
                primpairs[k]["incvector"].append(0)
    return primpairs


def checkCompNeo(pp1,pp2,k,e,mindg,maxdt,mindlen):
    """!
    Función que se encarga de revisar compatibilidad entre dos pares de partidores. Almacena las causas de incompatibilidad en un string (rson). Esta versión utiliza rangos de temperatura en vez de temperaturas discretas.
    \param <pp1> {Diccionario, tiene los datos de un par de partidores.}
    \param <pp2> {Diccionario, tiene los datos de un par de partidores.}
    \param <k> {Entero, posición asociada al primer par de partidores.}
    \param <e> {Entero, posición asociada al segundo par de partidores.}
    \param <mindg> {Numero, minimo deltaG tolerado.}
    \param <maxdt> {Numero, maxima diferencia de temperatura tolerada.}
    \param <mindlen> {Entero, minima diferencia de largo requerida.}
    """ 
    rson=[]
    comp=True
    f1=pp1["f"]
    r1=pp1["r"]
    f2=pp2["f"]
    r2=pp2["r"]
    gibf1f2=primer3.calcHeterodimer(f1,f2,mv_conc=gmvc).dg
    gibf1r2=primer3.calcHeterodimer(f1,r2,mv_conc=gmvc).dg
    gibr1f2=primer3.calcHeterodimer(r1,f2,mv_conc=gmvc).dg
    gibr1r2=primer3.calcHeterodimer(r1,r2,mv_conc=gmvc).dg
    rpe=pp2["relpos"]
    rpk=pp1["relpos"]
    if (pp1["matm"]>=pp2["matm"]):
        dt=pp1["tp"]-pp2["matm"]
    else:
        dt=pp2["tp"]-pp1["matm"]        
    dsize=abs(pp1["ampsize"]-pp2["ampsize"])
    if abs(rpe-rpk)==1:
        comp=False
        rson.append("Contiguedad")
    if gibf1f2<mindg:
        comp=False
        rson.append("F{} y F{} dan Gibbs {}".format(k,e,gibf1f2))
    if gibf1r2<mindg:
        comp=False
        rson.append("F{} y R{} dan Gibbs {}".format(k,e,gibf1r2))
    if gibr1f2<mindg:
        comp=False
        rson.append("R{} y F{} dan Gibbs {}".format(k,e,gibr1f2))
    if gibr1r2<mindg:
        comp=False
        rson.append("R{} y R{} dan Gibbs {}".format(k,e,gibr1r2))
    if dt>maxdt:
        comp=False
        rson.append("Diferencia de temperatura es {}".format(dt))
    if dsize<mindlen:
        comp=False
        rson.append("Diferencia de tamaño es {}".format(dsize))
    return rson


# In[15]:


def eqTemps(colorSet,primerDict,genseq):
    """!
    Función que intenta equiparar lo más posible las temperaturas dentro de un grupo.
    \param <colorSet> {Diccionario, grupos creados por el algoritmo de coloreo.}
    \param <primerDict> {Diccionario, diccionario de partidores.}
    \param <genseq> {Seq, secuencia del gen.}
    """ 
    primpairs=copy.deepcopy(primerDict)
    colors=copy.deepcopy(colorSet)
    print(colors)
    for c in colors:
        #print("activado color",c)
        maxogtemp=-9999
        templist=colors[c]["pairs"]
        for t in templist:
            temptm=primpairs[t]["tp"]
            if temptm > maxogtemp:
                maxogtemp=temptm
        for t in templist:
            exte=primpairs[t]["extend"]
            if (exte==0 and primpairs[t]["ftype"]!="gemini"):
                continue
            if (exte==0):
                posil=primpairs[t]["givstart"]
            #print("hi",t)
            tm=primpairs[t]["tp"]
            f=primpairs[t]["f"]
            r=primpairs[t]["r"]
            maxflen=primpairs[t]["maxflen"]
            maxrlen=primpairs[t]["maxrlen"]
            tmf=primer3.calcTm(f,mv_conc=gmvc)
            tmr=primer3.calcTm(r,mv_conc=gmvc)
            while tm<maxogtemp and (len(f)<30 or len(r)<30) and (len(f)<maxflen or len(r)<maxrlen):     
                if (tmf<=tm or len(r)==30) and len(f)<30 and len(f)<maxflen:
                    #print("se alargo f de {} en color {}".format(t,c))
                    if (exte==0):
                        f=str(extendLWG(genseq,f,posil,1))
                    else:
                        f=str(extendPrimer(genseq,f,"L",1))
                    tmf=primer3.calcTm(f,mv_conc=gmvc)
                else:
                    #print("se alargo r de {} en color {}".format(t,c))
                    if (exte==0):
                        re=extendRWGL(genseq,f,posil,r,1)
                    else:
                        _,re=extendTogether(genseq,f,r,1)
                    r=str(re)
                    tmr=primer3.calcTm(r,mv_conc=gmvc)
                tm=(tmf+tmr)/2.0
            primpairs[t]["r"]=r
            primpairs[t]["f"]=f
            primpairs[t]["tp"]=tm
        
    return primpairs


# In[16]:


def nxColoring(comppairs,mode=1):
    """!
    Función que realiza el algoritmo de coloreo utilizando una estrategía determinada.
    \param <comppairs> {Diccionario, diccionario de partidores incluyendo vector de compatibilidad.}
    \param <mode> {Entero, indica a la función que estrategia de coloreo debe usar.}
    """ 
    algos={
        1:'largest_first',
        2:'random_sequential',
        3:'smallest_last',
        4:'independent_set',
        5:'connected_sequential_bfs',
        6:'connected_sequential_dfs',
        7:'saturation_largest_first'
        
    }
    pprimpairs=copy.deepcopy(comppairs)
    network = nx.Graph()
    for k in pprimpairs:
        network.add_node(k)
        incom=pprimpairs[k]["incvector"]
        for i in range(len(incom)):
            if incom[i]==1:
                network.add_edge(k,i+1)


    #print(f"This network has now {network.number_of_nodes()} nodes.")
    #nx.draw_networkx(network, with_labels=True)
    d = nx.coloring.greedy_color(network, strategy=algos[mode])
    colors=set(d.values())
    colordic={}
    for c in colors:
        tempprimsets=[]
        for primset in d:
            if d[primset]==c:
                tempprimsets.append(primset)
        colordic[c+1]=tempprimsets

    #print(colordic)

    for c in colordic:
        #print("Color "+str(c)+":")
        for s in colordic[c]:
            #print(s," | ",pprimpairs[s]["tp"])
            None
        #print("----------------------------------------------------")
        None
    return colordic
            
def nxMinGroups(comppairs):
    """!
    Función que prueba distintas estrategias de coloreo para encontrar la que genera el menor numero de grupos.
    \param <comppairs> {Diccionario, diccionario de partidores incluyendo vector de compatibilidad.}
    """ 
    testpairs=copy.deepcopy(comppairs)
    colorModeResults={}
    for i in [1,3,4,5,6,7]:
        tempcodi=nxColoring(testpairs,i)
        nco=len(tempcodi)
        sumlen=0
        for c in tempcodi:
            sumlen+=len(tempcodi[c])
        avglen=sumlen/nco
        sumdif=0
        for c in tempcodi:
            sumdif+=abs(len(tempcodi[c])-avglen)
        tempd={"colorSet":tempcodi,"nColor":nco,"avgdif":sumdif}
        colorModeResults[i]=tempd
        #print (colorModeResults)
    for k in colorModeResults:
        print("Modo {} usa {} colores con diferencia de {}".format(k,colorModeResults[k]["nColor"]
                                                                   ,colorModeResults[k]["avgdif"]))
        None
    print("\n")
    lownum=99999
    dic2={}
    for k in colorModeResults:
        kcolors=colorModeResults[k]["nColor"]
        if kcolors<lownum:
            lownum=kcolors
    for k in colorModeResults:
        kcolors=colorModeResults[k]["nColor"]
        if kcolors==lownum:
            dic2[k]=colorModeResults[k]
    if len(dic2)>1:
        lownum2=99999
        dic3={}
        #print(list(dic2.keys()))
        for k in dic2:
            kdif=dic2[k]["avgdif"]
            if kdif<lownum2:
                lownum2=kdif
        #print(lownum2)
        for k in dic2:
            kdif=dic2[k]["avgdif"]
            #print(kdif,lownum2)
            if kdif==lownum2:
                #print(k,lownum2)
                dic3[k]=dic2[k]
        #print(dic3)
            #print(list(dic3.keys())[0])
        return dic3[list(dic3.keys())[0]]["colorSet"]
    else:
        print(dic2)
        #print(list(dic2.keys())[0])
        return dic2[list(dic2.keys())[0]]["colorSet"]
        
def colorSetAdapter(colorSet):
    todap=copy.deepcopy(colorSet)
    dictSol={}
    for k in todap:
        temp={"pairs":todap[k]}
        dictSol[k]=temp
    for cn in todap:
            print("Sets de partidores para color {}: \n".format(cn) + 
                  ', '.join(str(field) for field in dictSol[cn]["pairs"])+"\n")
    return dictSol
    
        


def write2cvs(colorSet,primerSet,newfilename="paneles_multiplex.csv"):
    """!
    Función encargada de escribir el panel multiplexado a un archivo ".csv".
    \param <colorSet> {Diccionario, grupos creados por algoritmo de coloreo.}
    \param <primerSet> {Diccionario, diccionario de partidores.}
    \param <newfilename> {String, nombre del archivo a escribir.}
    """ 
    primpairs=copy.deepcopy(primerSet)
    colorsets=copy.deepcopy(colorSet)
    sedico={}
    for p in primpairs:
        for k in colorsets:
            if p in colorsets[k]["pairs"]:
                sedico[p]=k
    primdictlist=[]
    for k in primpairs:
        setnum=k
        fw=primpairs[k]["f"]
        rv=primpairs[k]["r"]
        tm=primpairs[k]["tp"]
        grn=sedico[k]
        asp=primpairs[k]["ampsize"]
        filen=primpairs[k]["file"]
        tempdict={"SetNumber": setnum, 
                  "Forward": fw, "fTm": primer3.calcTm(fw,mv_conc=gmvc), "fGC": str(calcGC(fw)*100)+"%", "lenF":len(fw),
                  "Reverse": rv, "rTm": primer3.calcTm(rv,mv_conc=gmvc), "rGC": str(calcGC(rv)*100)+"%", "lenR":len(rv),
                  "AvgTm": tm, 'AmpliconSize': asp, "GroupNumber": grn, "File": filen}
        primdictlist.append(tempdict)

    dictfields = ['SetNumber', 'Forward', 'fTm', 'fGC', 'lenF',
                  'Reverse', 'rTm', 'rGC', 'lenR', 'AvgTm',
                  'AmpliconSize', 'GroupNumber', 'File']

    with open(newfilename, 'w') as csvfile: 
        writer = csv.DictWriter(csvfile, fieldnames = dictfields) 
        writer.writeheader() 
        writer.writerows(primdictlist) 
        print("\nArchivo {} creado en directorio.".format(newfilename))


# In[17]:

#FastaFixData
def checkExtend(fixList):
    """!
    Función encargada de evaluar si es posible extender un partidor.
    \param <fixList> {Lista, partidores que fueron procesados por función "FastaFixData".}
 
    """ 
    for f in fixList:
        for k in f:
            giv=f[k]["ampsize"]
            calc=f[k]["calcsize"]
            ext=f[k]["extend"]
            #print(giv,calc,calc-giv)
            print("Par {}{}se puede extender.".format(k,(" *NO* " if ext==0 else " ")))
        print("-----------------------------------------------------------------------------------------------")

def fixPositions(tepairs,teseq,sfnames):
    primpairs=copy.deepcopy(tepairs)
    seqlens=[len(Word2Seq(s)) for s in sfnames]
    foundfiles=[]
    for k in primpairs:
        if primpairs[k]["file"] not in foundfiles:
            foundfiles.append(primpairs[k]["file"])
    modsdict={}
    for i,f in enumerate(foundfiles):
        modsdict[f]=sum(seqlens[:i])
        
    for k in primpairs:
        primpairs[k]["givstart"]+=modsdict[primpairs[k]["file"]]
        
    return primpairs

# In[18]:


#main function
#valores ngs
#mv_conc=50
#dv_conc=3
#dg=-10000

def multiplexPrimers(sfname,sftype,pfname,pftype,isMultiple=False,mindg=-13627,maxdt=5,mindlen=0,maxcolor=100,maxlen=30,mv_conc=164):
    """!
    Función para generar paneles de partidores multiplexados.
    \param <sfname> {String, nombre del archivo de secuencia del gen.}
    \param <sftype> {String, tipo del archivo de secuencia (por ejemplo, "word").}
    \param <pfname> {String, nombre del archivo de partidores.}
    \param <pftype> {String, tipo del archivo de partidores (por ejemplo, "GEMIN1").}
    \param <isMultiple> {Booleano, si es un panel multi-gen o no.}
    \param <mindg> {Numero, minimo deltaG tolerado.}
    \param <maxdt> {Numero, maxima diferencia de temperatura tolerada.}
    \param <mindlen> {Entero, minima diferencia de largo requerida.}
    \param <maxcolor> {Entero, maximo numero de colores.}
    \param <maxlen> {Entero, maximo largo para partidores.}
    \param <mv_conc> {Numero, concentración de iones monovalentes en medio de reacción.}
    """ 
    global gmvc
    gmvc=mv_conc
    if isMultiple:
        firstname="multiplexed_"+pfname[0]
        pl=[]
        fl=[]
        for s in pfname:
            pl.append(s)
        pairs=loadMultiple(pl,pftype)
        for s in sfname:
            ts=Word2Seq(s)
            fl.append(ts)
        seq=fl[0]
        for i in range(1,len(fl)):
            seq+=fl[i]         
        pairs=fixPositions(pairs,seq,sfname)        
          
    else:
        firstname="multiplexed_"+pfname
        pairs=loadPrimPairs(pfname,pftype)
        seq=Word2Seq(sfname)
        
    seqfix=FastaFixData(pairs,seq)
    checkExtend([seqfix])
    deltaTp=findMaxTemp(seqfix,seq,maxlen)
    
    print("\nCREANDO PRIMER GRAFICO...\n")
    
    neocomp=checkGroupCompNeo(deltaTp,mindg,maxdt,mindlen)
    writeReasons(neocomp,firstname[:-4]+"_first_graph_reasons.txt")
    newColors=colorSetAdapter(nxMinGroups(neocomp))
    
    print("\nAJUSTANDO TM...\n")
    
    balancedp=eqTemps(newColors,deltaTp,seq)
    
    print("\nCREANDO SEGUNDO GRAFICO...\n")
    
    comppairs2=checkGroupComp(balancedp,mindg,maxdt,mindlen)
    writeReasons(comppairs2,firstname[:-4]+"_second_graph_reasons.txt")
    colorSets2=colorSetAdapter(nxMinGroups(comppairs2))
    if __name__ == "__main__":
        write2cvs(colorSets2,balancedp,newfilename=firstname)
    else:
        return colorSets2,balancedp


if __name__ == "__main__":
    
    import argparse
    parser = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument('-gfn', '--genfilename', action='append', help='Nombre del archivo de secuencia del gen (Repetir para multiple).', required=True, default=argparse.SUPPRESS)
    parser.add_argument('-pfn', '--primerfilename', action='append', help='Nombre del archivo de partidores (Repetir para multiple).', required=True, default=argparse.SUPPRESS)
    parser.add_argument('-mdg', '--mindg',help='Minimo deltaG tolerado.',type=int ,default=-7000)
    parser.add_argument('-mdt', '--maxdt',help='Maxima diferencia de temperatura tolerada.',type=int ,default=5)
    parser.add_argument('-mdl', '--mindlen',help='Minima diferencia de largo requerida.',type=int ,default=0)
    parser.add_argument('-ml', '--maxlen',help='Maximo largo para partidores.',type=int ,default=30)
    parser.add_argument('-mvc', '--mvconc',help='Concentración de iones monovalentes en medio de reacción.',type=int ,default=50)

    args = parser.parse_args()
    dargs = vars(args)
    
    if len(dargs["genfilename"])!=len(dargs["primerfilename"]):
        raise TypeError("Debe existir una secuencia por cada panel.") 
    
    argseq=dargs["genfilename"]
    argprim=dargs["primerfilename"]
    multivar=len(dargs["primerfilename"])>1
    if not(multivar):
        argseq=argseq[0]
        argprim=argprim[0]
    
    multiplexPrimers(argseq,"word",argprim,"GEMIN1",mv_conc=dargs["mvconc"],mindg=dargs["mindg"],maxdt=dargs["maxdt"], mindlen=dargs["mindlen"],maxlen=dargs["maxlen"],isMultiple=multivar)




